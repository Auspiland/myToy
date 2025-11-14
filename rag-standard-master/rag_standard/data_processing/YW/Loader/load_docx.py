import os, re, sys
from docx import Document

BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '../../../..')) # 상대경로로 root 설정 (Rag_Standard_v2)
sys.path.append(BASE_DIR)
from rag_standard.utils import setup_logger

global log
log = setup_logger.setup_logger('LoadDOCX', 'load_docx')

class LoadDOCX:
    """ load_docx만 호출하면 됨 """
    def __init__(self):
        pass

    def load_docx(self, file_path):
        """
        워드 문서(.docx) 파일을 로드하고 내용을 추출하는 메소드 >> Python-docx 사용
        
        Args:
            file_path: 워드 문서 파일 경로
            
        Returns:
            tuple: (파일명, 확장자, 수정일, 생성일, 텍스트만 추출한 딕셔너리, 텍스트와 테이블을 모두 추출한 딕셔너리, 전체 텍스트를 저장한 문자열)

        >> 참고: DOCX는 전체 문서를 하나의 청크로 보고 내용을 추출하기 때문에, 아래와 같은 딕셔너리 형식으로 추출됨
            {
                '1': 문서 내용
            }
        """
        file_name_without_ext, extension = os.path.basename(file_path).rsplit('.', 1) # 파일명과 확장자 분리

        try:
            doc = Document(file_path) # DOCX 파일 열기

            # 문서의 수정일자와 생성일자 추출 - 메타데이터가 없는 경우 파일 시스템의 날짜 정보 사용
            modified_date = doc.core_properties.modified.strftime("%Y-%m-%d %H:%M:%S") if hasattr(doc, 'core_properties') and doc.core_properties.modified else str(os.path.getmtime(file_path).strftime("%Y-%m-%d %H:%M:%S"))
            created_date = doc.core_properties.created.strftime("%Y-%m-%d %H:%M:%S") if hasattr(doc, 'core_properties') and doc.core_properties.created else str(os.path.getctime(file_path).strftime("%Y-%m-%d %H:%M:%S"))
            
            # 모든 단락의 텍스트 추출
            full_text = [para.text for para in doc.paragraphs]
            txt = '\n'.join(full_text)
            
            # 표 데이터 추출 (일반 텍스트 형식)
            tables_data = []
            for table in doc.tables:
                if len(table.rows) > 0:
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        tables_data.append('\n'.join(row_data))
            
            tbl = '\n'.join(tables_data)

            # 표 데이터를 마크다운 형식으로 변환
            markdown_tables = []
            for table in doc.tables:
                markdown_table = []
                # 헤더 행 처리
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                markdown_table.append('| ' + ' | '.join(headers) + ' |')
                markdown_table.append('|' + ' --- |' * len(headers))
                
                # 데이터 행 처리
                for row in table.rows[1:]: # 헤더 행 제외
                    row_data = [cell.text.strip() for cell in row.cells]
                    markdown_table.append('| ' + ' | '.join(row_data) + ' |')
                
                markdown_tables.append('\n'.join(markdown_table))
            mark = '\n\n'.join(markdown_tables)
            
            # 텍스트와 표 데이터 결합
            combined_text = txt + '\n' + tbl # 텍스트만 있는 결과
            combined_table = txt + '\n\n' + mark # 마크다운 형식의 표 포함

            # 공백 정리 (연속된 공백과 줄바꿈 제거)
            text_only_result = re.sub(r'\n{2,}', '\n', re.sub(r' {2,}', ' ', combined_text))
            text_and_table_result = re.sub(r'\n{3,}', '\n', re.sub(r' {2,}', ' ', combined_table))

            log.info(f"Completed loading docx file: {os.path.basename(file_path)}")

        except Exception as e:
            log.error(f"Cannot load DOCX file: {file_path} due to {e}")
            return True

        # 결과 반환: (파일명, 확장자, 수정일, 생성일, 텍스트 결과, 표+텍스트 결과) - 텍스트 결과와 표+텍스트 결과는 각각 딕셔너리 형태로 반환 (키: '1')
        return (file_name_without_ext, extension.lower(), modified_date, created_date, {'1':text_only_result}, {'1':text_and_table_result}, text_and_table_result)